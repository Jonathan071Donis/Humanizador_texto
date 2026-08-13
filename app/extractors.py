"""
Extracts raw text content from uploaded files, entirely in memory
(io.BytesIO) - nothing is ever written to disk.
"""
from __future__ import annotations

import io
from typing import Tuple

TEXT_EXTENSIONS = {
    ".txt", ".md", ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".cpp",
    ".c", ".h", ".hpp", ".go", ".rs", ".php", ".rb", ".cs", ".swift",
    ".kt", ".sh", ".sql", ".json", ".yaml", ".yml", ".html", ".css",
    ".xml", ".ini", ".cfg", ".toml",
}


class ExtractionError(Exception):
    pass


def _ext(filename: str) -> str:
    idx = filename.rfind(".")
    return filename[idx:].lower() if idx != -1 else ""


def extract_pdf(data: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        pdfplumber = None

    if pdfplumber is not None:
        try:
            text_parts = []
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts)
        except Exception:
            pass  # fall back to PyPDF2 below

    try:
        from PyPDF2 import PdfReader
    except ImportError as e:
        raise ExtractionError("No PDF library available (pdfplumber / PyPDF2)") from e

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as e:
        raise ExtractionError("python-docx is not installed") from e

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_content(filename: str, data: bytes) -> Tuple[str, str]:
    """Returns (extracted_text, file_type_label)."""
    ext = _ext(filename)

    if ext == ".pdf":
        return extract_pdf(data), "pdf"
    if ext == ".docx":
        return extract_docx(data), "docx"
    if ext in TEXT_EXTENSIONS or ext == "":
        return extract_text_bytes(data), ext.lstrip(".") or "text"

    # Unknown extension: try to treat as text; if it looks binary, refuse.
    try:
        text = extract_text_bytes(data)
        return text, ext.lstrip(".")
    except Exception as e:
        raise ExtractionError(f"Unsupported file type: {ext}") from e
